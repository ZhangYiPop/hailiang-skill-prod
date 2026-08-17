from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = Path(__file__).resolve().parents[1] / "docs" / "大模型产品外部测试API接口规范.docx"


def shade(cell, fill: str) -> None:
    props = cell._tc.get_or_add_tcPr()
    element = OxmlElement("w:shd")
    element.set(qn("w:fill"), fill)
    props.append(element)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120) -> None:
    props = cell._tc.get_or_add_tcPr()
    margins = props.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        props.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_widths(table, widths_cm: list[float]) -> None:
    for row in table.rows:
        for cell, width in zip(row.cells, widths_cm):
            cell.width = Cm(width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    normal = doc.styles["Normal"]
    normal.font.name = "Songti SC"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Songti SC")
    normal.font.size = Pt(10.5)
    for name, size, color in (("Title", 20, "1F4E79"), ("Heading 1", 15, "1F4E79"), ("Heading 2", 12, "2F75B5")):
        style = doc.styles[name]
        style.font.name = "Heiti SC"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Heiti SC")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)


def add_code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
    run.font.size = Pt(9)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, value in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = value
        shade(cell, "D9EAF7")
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for row_values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row_values):
            cell.text = value
    set_widths(table, widths)
    doc.add_paragraph()


def main() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_doc(doc)
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("大模型产品外部测试 API 接口规范")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("面向外部测评与联调调用 | 版本：v1").italic = True

    doc.add_heading("一、接口规范要求", level=1)
    doc.add_paragraph("本接口用于外部人员测试现有大模型对话能力。每次调用自动创建新的用户、档案、会话和运行实例，调用方无需管理历史会话。调用方负责在 dialogue 中拼接完整上下文，服务端使用最后一条 user 消息作为当前问题。")
    doc.add_heading("1.1 请求规范", level=2)
    doc.add_paragraph("请求接口必须使用 JSON 格式，并通过固定 API Key 鉴权。API Key 由部署环境变量 HAILIANG_EXTERNAL_API_KEY 配置。")
    add_code(doc, "Content-Type: application/json\nAuthorization: Bearer ${api_key}")
    doc.add_heading("1.2 请求体", level=2)
    add_code(doc, '{\n  "model": "default",\n  "max_tokens": 1024,\n  "stream": false,\n  "dialogue": [\n    {"role": "user", "content": "请介绍一下北京。"},\n    {"role": "model", "content": "北京是中国的首都。"},\n    {"role": "user", "content": "请再介绍当地美食。"}\n  ]\n}')
    add_table(doc, ["字段", "类型", "必填", "默认值", "说明"], [
        ["model", "字符串", "否", "default", "模型标识；当前使用服务默认模型。"],
        ["max_tokens", "整数", "否", "1024", "输出长度上限，范围 1-32768。"],
        ["stream", "布尔", "否", "false", "false 返回 JSON，true 返回 SSE。"],
        ["dialogue", "数组", "是", "无", "至少一条消息，最后一条必须是 user。"],
        ["dialogue[].role", "user/model/assistant", "是", "无", "model 在内部转换为 assistant。"],
        ["dialogue[].content", "非空字符串", "是", "无", "消息内容。"],
    ], [2.8, 3.0, 1.4, 2.0, 7.2])
    doc.add_paragraph("注意：model 和 max_tokens 为兼容评测平台保留的请求字段；当前服务使用部署配置中的默认模型参数，不支持通过请求动态切换模型。")

    doc.add_heading("二、响应规范", level=1)
    doc.add_paragraph("非流式请求返回 HTTP 200，业务成功或失败均在 JSON 最外层返回固定字段。")
    add_code(doc, '{\n  "content": "模型回答内容",\n  "choices": [],\n  "status": "success",\n  "reason": "success",\n  "session_id": "sess_external_xxx",\n  "request_id": "req_xxx"\n}')
    doc.add_paragraph("业务执行失败时 content 为空、status 为 failed，reason 返回失败原因；API Key 错误、请求格式错误和限流等建连前错误使用 HTTP 4xx/429。")
    doc.add_heading("2.1 流式响应", level=2)
    doc.add_paragraph("stream=true 时返回 text/event-stream。每个 data 事件只包含外部协议需要的增量文本，增量内容放在 choices[].delta；最后一个事件包含 finish_reason=stop。内部会话状态、Facts、Skill 状态和调试字段不会对外暴露。")
    add_code(doc, 'data: {"content":"","choices":[{"delta":"北京"}],"status":"success","reason":"success"}\n\ndata: {"content":"","choices":[{"delta":"","finish_reason":"stop"}],"status":"success","reason":"success","session_id":"sess_external_xxx"}')

    doc.add_heading("三、API 接口信息", level=1)
    add_table(doc, ["项目", "内容"], [
        ["接口名称", "外部大模型测试接口"],
        ["请求方式", "POST"],
        ["请求地址", "/api/v1/external/chat"],
        ["鉴权方式", "Authorization: Bearer ${api_key}"],
        ["默认响应", "非流式 JSON"],
        ["会话策略", "每次调用自动新建，不复用历史会话"],
        ["上下文策略", "调用方传入完整 dialogue，最后一条 user 为当前问题"],
    ], [3.5, 12.9])
    doc.add_heading("3.1 非流式调用示例", level=2)
    add_code(doc, 'curl "$ALGORITHM_BASE/api/v1/external/chat" \\\n+  -H "Authorization: Bearer ${HAILIANG_EXTERNAL_API_KEY}" \\\n+  -H "Content-Type: application/json" \\\n+  --data-raw \'{"dialogue":[{"role":"user","content":"请介绍北京。"}]}\'')
    doc.add_heading("3.2 流式调用示例", level=2)
    add_code(doc, 'curl -N "$ALGORITHM_BASE/api/v1/external/chat" \\\n+  -H "Authorization: Bearer ${HAILIANG_EXTERNAL_API_KEY}" \\\n+  -H "Content-Type: application/json" \\\n+  --data-raw \'{"stream":true,"dialogue":[{"role":"user","content":"请介绍北京。"}]}\'')
    doc.add_heading("3.3 错误码", level=2)
    add_table(doc, ["HTTP 状态", "错误码", "说明"], [
        ["401", "INVALID_API_KEY", "Authorization 缺失或 API Key 不正确。"],
        ["422", "REQUEST_VALIDATION_ERROR", "请求字段、类型或内容不符合规范。"],
        ["422", "DIALOGUE_LAST_MESSAGE_MUST_BE_USER", "dialogue 最后一条消息不是 user。"],
        ["429", "LLM_RATE_LIMITED", "模型服务限流。"],
        ["503", "EXTERNAL_API_NOT_CONFIGURED", "服务端未配置外部接口 API Key。"],
    ], [2.5, 5.2, 8.7])
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
